#!/usr/bin/env python3
"""
Advanced Bajaj Finserv Customer Portal Data Ingestion and Knowledge Graph Creation
This script crawls and extracts structured/unstructured content from Bajaj Finserv's customer portal,
including financial product data and downloadable files (PDF, DOCX, XLSX, CSV), and creates a comprehensive knowledge graph.
"""

import asyncio
import aiohttp
import json
import csv
import time
import logging
import re
import os
from pathlib import Path
from urllib.parse import urljoin, urlparse
from typing import Dict, List, Set, Tuple, Optional
from dataclasses import dataclass
from datetime import datetime
import sqlite3
from dotenv import load_dotenv

# Web scraping and parsing
from bs4 import BeautifulSoup
import lxml

# NLP and Knowledge Graph
import spacy
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.tag import pos_tag
from nltk import ne_chunk

# Document processing
import PyPDF2
import openpyxl
from docx import Document
import mammoth

# Graph and visualization
import networkx as nx
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from collections import defaultdict, Counter

# Load environment variables
load_dotenv()
DECAY_FACTOR = float(os.getenv('DECAY_FACTOR', 0.99))

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class WebContent:
    """Data class for web content"""
    url: str
    title: str
    content: str
    content_type: str
    meta_tags: Dict[str, str]
    links: List[str]
    tables: List[Dict]
    accessibility_tags: Dict[str, List[str]]
    timestamp: datetime
    depth: int
    file_path: Optional[Path] = None

@dataclass
class Entity:
    """Data class for extracted entities"""
    text: str
    label: str
    start: int
    end: int
    confidence: float = 0.0

@dataclass
class Relationship:
    """Data class for entity relationships"""
    source: str
    target: str
    relation: str
    confidence: float = 0.0

class BajajFinservCrawler:
    """Advanced web crawler for Bajaj Finserv with depth control, prioritization, and file handling"""
    
    def __init__(self, base_url: str = os.getenv('BASE_URL', "https://www.bajajfinservhealth.in/"), 
                 max_pages: int = int(os.getenv('MAX_PAGES', 10000)), 
                 max_depth: int = int(os.getenv('MAX_DEPTH', 10)), 
                 delay: float = float(os.getenv('DELAY', 0.1)), 
                 output_dir: Path = Path(os.getenv('OUTPUT_DIR', "bajajfinserv_data"))):
        self.base_url = base_url
        self.max_pages = max_pages
        self.max_depth = max_depth
        self.delay = delay
        self.output_dir = output_dir
        self.visited_urls: Set[str] = set()
        self.scraped_content: List[WebContent] = []
        self.url_queue = []  # (depth, url)
        self.session = aiohttp.ClientSession()
        self.headers = {
            'User-Agent': os.getenv('USER_AGENT', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
        }
        self.priority_keywords = os.getenv('PRIORITY_KEYWORDS', '').split(', ')
        self.doc_processor = DocumentProcessor()
    
    def is_valid_url(self, url: str) -> bool:
        """Check if URL is valid and within domain"""
        parsed = urlparse(url)
        return (
            parsed.netloc.endswith('bajajfinserv.in') and
            url not in self.visited_urls
        )
    
    async def download_file(self, url: str, response: aiohttp.ClientResponse) -> Optional[Path]:
        """Download file and save to disk"""
        try:
            content_type = response.headers.get('Content-Type', '').lower()
            if 'pdf' in content_type:
                ext = '.pdf'
            elif 'excel' in content_type or 'spreadsheet' in content_type:
                ext = '.xlsx'
            elif 'word' in content_type:
                ext = '.docx'
            elif 'csv' in content_type:
                ext = '.csv'
            else:
                ext = '.bin'  # Default for unknown types
            filename = url.split('/')[-1] or f"file_{int(time.time())}{ext}"
            save_path = self.output_dir / 'downloads' / filename
            save_path.parent.mkdir(exist_ok=True)
            with open(save_path, 'wb') as f:
                async for chunk in response.content.iter_chunked(1024):
                    f.write(chunk)
            return save_path
        except Exception as e:
            logger.error(f"Error downloading {url}: {str(e)}")
            return None
    
    async def scrape_page(self, url: str, depth: int) -> Optional[WebContent]:
        """Scrape a single page or file asynchronously"""
        try:
            async with self.session.get(url, headers=self.headers, timeout=int(os.getenv('TIMEOUT', 10))) as response:
                if response.status != 200:
                    return None
                content_type = response.headers.get('Content-Type', '').lower()
                
                if 'text/html' in content_type:
                    # Process as HTML
                    html = await response.text()
                    soup = BeautifulSoup(html, 'lxml')
                    title = soup.find('title')
                    title_text = title.get_text(strip=True) if title else ""
                    
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    content = soup.get_text(separator=' ', strip=True)
                    links = [urljoin(url, link['href']) for link in soup.find_all('a', href=True) 
                             if self.is_valid_url(urljoin(url, link['href']))]
                    tables = self.extract_tables(soup)
                    
                    web_content = WebContent(
                        url=url,
                        title=title_text,
                        content=content,
                        content_type=content_type,
                        meta_tags=self.extract_meta_tags(soup),
                        links=links,
                        tables=tables,
                        accessibility_tags=self.extract_accessibility_tags(soup),
                        timestamp=datetime.now(),
                        depth=depth
                    )
                else:
                    # Handle as downloadable file
                    file_path = await self.download_file(url, response)
                    if not file_path:
                        return None
                    
                    # Extract text based on file type
                    if file_path.suffix == '.pdf':
                        content = self.doc_processor.process_pdf(str(file_path))
                    elif file_path.suffix == '.docx':
                        content = self.doc_processor.process_docx(str(file_path))
                    elif file_path.suffix == '.xlsx':
                        content = self.doc_processor.process_xlsx(str(file_path))
                    elif file_path.suffix == '.csv':
                        content = self.doc_processor.process_csv(str(file_path))
                    else:
                        content = ""  # Unsupported file type
                    
                    web_content = WebContent(
                        url=url,
                        title=file_path.name,
                        content=content,
                        content_type=content_type,
                        meta_tags={},
                        links=[],
                        tables=[],
                        accessibility_tags={},
                        timestamp=datetime.now(),
                        depth=depth,
                        file_path=file_path
                    )
                
                logger.info(f"Scraped: {url} at depth {depth} ({content_type})")
                return web_content
        except Exception as e:
            logger.error(f"Error scraping {url}: {str(e)}")
            return None
    
    def extract_meta_tags(self, soup: BeautifulSoup) -> Dict[str, str]:
        """Extract meta tags from HTML"""
        meta_tags = {}
        for tag in soup.find_all('meta'):
            name = tag.get('name') or tag.get('property')
            content = tag.get('content')
            if name and content:
                meta_tags[name] = content
        return meta_tags
    
    def extract_tables(self, soup: BeautifulSoup) -> List[Dict]:
        """Extract table data from HTML"""
        tables = []
        for table in soup.find_all('table'):
            table_data = {'headers': [], 'rows': []}
            headers = table.find_all('th')
            if headers:
                table_data['headers'] = [th.get_text(strip=True) for th in headers]
            for row in table.find_all('tr'):
                cells = row.find_all(['td', 'th'])
                if cells:
                    table_data['rows'].append([cell.get_text(strip=True) for cell in cells])
            if table_data['rows']:
                tables.append(table_data)
        return tables
    
    def extract_accessibility_tags(self, soup: BeautifulSoup) -> Dict[str, List[str]]:
        """Extract accessibility-related tags"""
        accessibility = {'aria_labels': [], 'alt_texts': [], 'role_attributes': []}
        for element521 in soup.find_all(attrs={'aria-label': True}):
            accessibility['aria_labels'].append(element521.get('aria-label'))
        for img in soup.find_all('img', alt=True):
            accessibility['alt_texts'].append(img.get('alt'))
        for element521 in soup.find_all(attrs={'role': True}):
            accessibility['role_attributes'].append(element521.get('role'))
        return accessibility
    
    async def crawl(self, start_urls: List[str] = None):
        """Main crawling method with depth control and prioritization"""
        if start_urls is None:
            start_urls = [self.base_url]
        
        for url in start_urls:
            self.url_queue.append((0, url))
        
        while self.url_queue and len(self.scraped_content) < self.max_pages:
            self.url_queue.sort(key=lambda x: (x[0], -self.get_url_priority(x[1])))
            current_depth, current_url = self.url_queue.pop(0)
            
            if current_url in self.visited_urls or current_depth > self.max_depth:
                continue
            
            self.visited_urls.add(current_url)
            web_content = await self.scrape_page(current_url, current_depth)
            if web_content:
                self.scraped_content.append(web_content)
                for link in web_content.links:
                    if link not in self.visited_urls:
                        self.url_queue.append((current_depth + 1, link))
            await asyncio.sleep(self.delay)
        
        logger.info(f"Crawling completed. Scraped {len(self.scraped_content)} pages.")
        return self.scraped_content
    
    def get_url_priority(self, url: str) -> int:
        """Calculate priority based on keywords"""
        url_lower = url.lower()
        return sum(keyword in url_lower for keyword in self.priority_keywords)
    
    async def close(self):
        """Close the aiohttp session"""
        await self.session.close()

class DocumentProcessor:
    """Process various document types"""
    
    @staticmethod
    def process_pdf(file_path: str) -> str:
        """Extract text from PDF"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                return "".join(page.extract_text() for page in reader.pages)
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def process_docx(file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def process_xlsx(file_path: str) -> str:
        """Extract text from XLSX"""
        try:
            workbook = openpyxl.load_workbook(file_path)
            text = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join(str(cell) for cell in row if cell is not None) + "\n"
            return text
        except Exception as e:
            logger.error(f"Error processing XLSX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def process_csv(file_path: str) -> str:
        """Extract text from CSV"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                text = ""
                for row in reader:
                    text += " ".join(row) + "\n"
            return text
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {str(e)}")
            return ""

class NLPProcessor:
    """Advanced NLP processing for entity and relationship extraction"""
    
    def __init__(self):
        nltk.download('punkt', quiet=True)
        nltk.download('averaged_perceptron_tagger', quiet=True)
        nltk.download('maxent_ne_chunker', quiet=True)
        nltk.download('words', quiet=True)
        nltk.download('stopwords', quiet=True)
        
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Install: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        self.stop_words = set(stopwords.words('english'))
        self.domain_entities = {
            'FINANCIAL_PRODUCT': ['loan', 'fixed deposit', 'insurance', 'card'],
            'DOCUMENT': ['statement of account', 'interest certificate', 'No Dues Certificate'],
            'SERVICE': ['EMI payment', 'part-prepayment', 'foreclosure', 'profile update'],
            'CONTACT_INFO': ['mobile number', 'email ID', 'PAN', 'date of birth']
        }
    
    def extract_entities(self, text: str) -> List[Entity]:
        """Extract entities using spaCy and domain-specific rules"""
        entities = []
        if self.nlp:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ in ['DATE', 'TIME']:
                    entities.append(Entity(ent.text, 'TEMPORAL', ent.start_char, ent.end_char, 0.8))
                else:
                    entities.append(Entity(ent.text, ent.label_, ent.start_char, ent.end_char, 0.8))
        
        for label, entity_list in self.domain_entities.items():
            for entity in entity_list:
                if entity in text:
                    start = text.find(entity)
                    entities.append(Entity(entity, label, start, start + len(entity), 0.9))
        return entities
    
    def extract_relationships(self, text: str, entities: List[Entity]) -> List[Relationship]:
        """Extract relationships using dependency parsing and temporal associations"""
        relationships = []
        if self.nlp:
            doc = self.nlp(text)
            for sent in doc.sents:
                for token in sent:
                    if token.dep_ in ('attr', 'dobj', 'pobj'):
                        subject = [w for w in token.head.lefts if w.dep_ == 'nsubj']
                        if subject:
                            relationships.append(Relationship(
                                subject[0].text, token.text, token.dep_, 0.7
                            ))
                    # Check for temporal relationships
                    if token.ent_type_ == 'DATE':
                        for entity in entities:
                            if entity.label != 'TEMPORAL' and entity.start <= token.idx <= entity.end:
                                relationships.append(Relationship(
                                    entity.text, token.text, 'ASSOCIATED_WITH', 0.7
                                ))
        return relationships
    
    def extract_attributes(self, text: str, entities: List[Entity]) -> Dict[str, Dict[str, str]]:
        """Extract attribute values using regex (e.g., EMI amounts)"""
        attributes = {}
        for entity in entities:
            if entity.label == 'FINANCIAL_PRODUCT':
                # Pattern for "EMI amount is Rs. X" or similar
                pattern = re.compile(r'EMI amount is Rs\. (\d+)', re.IGNORECASE)
                match = pattern.search(text)
                if match:
                    attributes[entity.text] = {'emi_amount': match.group(1)}
        return attributes

class KnowledgeGraph:
    def __init__(self):
        self.entities = {}
        self.graph = nx.DiGraph()
        self.ontology = {
            'FINANCIAL_PRODUCT': {'attributes': ['type', 'status']},
            'DOCUMENT': {'attributes': ['name', 'date']},
            'SERVICE': {'attributes': ['type', 'status']},
            'ORGANIZATION': {'attributes': ['name', 'location']},
            'TEMPORAL': {'attributes': ['date', 'time']}
        }

    def add_entity(self, entity: Entity):
        """Add entity with ontology attributes and confidence decay"""
        entity_id = f"{entity.label}_{entity.text}"
        if entity_id not in self.entities:
            self.entities[entity_id] = {
                'text': entity.text,
                'label': entity.label,
                'confidence': entity.confidence,
                'count': 1,
                'attributes': {}  # Initialize as an empty dictionary
            }
            self.graph.add_node(entity_id, **self.entities[entity_id])
        else:
            self.entities[entity_id]['count'] += 1
            decay = DECAY_FACTOR ** (self.entities[entity_id]['count'] - 1)
            new_confidence = self.entities[entity_id]['confidence'] + (entity.confidence * decay)
            self.entities[entity_id]['confidence'] = min(new_confidence, 1.0)
    
    def add_relationship(self, relationship: Relationship):
        """Add relationship to the graph"""
        source_id = self.find_entity_id(relationship.source)
        target_id = self.find_entity_id(relationship.target)
        if source_id and target_id:
            self.graph.add_edge(source_id, target_id, 
                              relation=relationship.relation, 
                              confidence=relationship.confidence)
    
    def find_entity_id(self, entity_text: str) -> Optional[str]:
        """Find entity ID by text"""
        for entity_id, data in self.entities.items():
            if data['text'].lower() == entity_text.lower():
                return entity_id
        return None
    
    def get_statistics(self) -> Dict:
        """Get graph statistics"""
        return {
            'num_entities': len(self.entities),
            'num_relationships': self.graph.number_of_edges(),
            'num_nodes': self.graph.number_of_nodes(),
            'num_edges': self.graph.number_of_edges()
        }
    
    def export_to_json(self, filename: str):
        """Export graph to JSON"""
        data = {
            'entities': self.entities,
            'relationships': [
                {
                    'source': self.entities[source]['text'],
                    'target': self.entities[target]['text'],
                    'relation': data['relation'],
                    'confidence': data['confidence']
                }
                for source, target, data in self.graph.edges(data=True)
            ],
            'statistics': self.get_statistics()
        }
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
    
    def visualize(self, filename: str = 'knowledge_graph.png'):
        """Visualize the graph with Noto Sans font for better character support"""
        font_path = os.getenv('FONT_PATH', r'assets\NotoSans-Regular.ttf')
        if not os.path.exists(font_path):
            raise FileNotFoundError(f"Font file not found: {font_path}")
        
        font_prop = fm.FontProperties(fname=font_path)
        fm.fontManager.addfont(font_path)
        font_name = font_prop.get_name()
        plt.rcParams['font.family'] = font_name
        
        plt.figure(figsize=(15, 10))
        pos = nx.spring_layout(self.graph, k=1, iterations=50)
        
        entity_types = set(data['label'] for data in self.entities.values())
        colors = plt.cm.Set3(range(len(entity_types)))
        color_map = dict(zip(entity_types, colors))
        node_colors = [color_map[self.entities[node]['label']] for node in self.graph.nodes()]
        
        nx.draw_networkx_nodes(self.graph, pos, node_color=node_colors, node_size=300, alpha=0.8)
        nx.draw_networkx_edges(self.graph, pos, alpha=0.5, arrows=True)
        
        labels = {node: self.entities[node]['text'][:10] for node in self.graph.nodes()}
        nx.draw_networkx_labels(self.graph, pos, labels, font_size=8, font_family=font_name)
        
        plt.title("Bajaj Finserv Knowledge Graph", fontfamily=font_name)
        plt.axis('off')
        plt.savefig(filename, dpi=300)
        plt.close()

class BajajFinservKnowledgeExtractor:
    """Main class for Bajaj Finserv knowledge extraction"""
    
    def __init__(self, output_dir: str = os.getenv('OUTPUT_DIR', "bajajfinserv_data")):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.crawler = BajajFinservCrawler(output_dir=self.output_dir)
        self.nlp_processor = NLPProcessor()
        self.knowledge_graph = KnowledgeGraph()
        self.db_conn = sqlite3.connect(self.output_dir / os.getenv('DATABASE_PATH', 'bajajfinserv.db'))
        self.create_db_tables()
    
    def create_db_tables(self):
        """Create SQLite tables"""
        cursor = self.db_conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS web_content (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            url TEXT UNIQUE,
            title TEXT,
            content TEXT,
            content_type TEXT,
            meta_tags TEXT,
            links TEXT,
            tables TEXT,
            accessibility_tags TEXT,
            timestamp TEXT,
            depth INTEGER,
            file_path TEXT
        )''')
        self.db_conn.commit()
    
    async def run_data_ingestion(self):
        """Run data ingestion process"""
        logger.info("Starting Bajaj Finserv data ingestion...")
        start_urls = [
            "https://www.bajajfinserv.in/",
            "https://www.bajajfinserv.in/loans",
            "https://www.bajajfinserv.in/fixed-deposit",
            "https://www.bajajfinserv.in/insurance",
            "https://www.bajajfinserv.in/customer-portal"
        ]
        scraped_content = await self.crawler.crawl(start_urls)
        self.save_scraped_content(scraped_content)
        return scraped_content
    
    def save_scraped_content(self, content: List[WebContent]):
        """Save scraped content to database"""
        cursor = self.db_conn.cursor()
        for item in content:
            cursor.execute('''INSERT OR IGNORE INTO web_content (
                url, title, content, content_type, meta_tags, links, tables, 
                accessibility_tags, timestamp, depth, file_path
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''', (
                item.url, item.title, item.content, item.content_type,
                json.dumps(item.meta_tags), json.dumps(item.links), json.dumps(item.tables),
                json.dumps(item.accessibility_tags), item.timestamp.isoformat(), item.depth,
                str(item.file_path) if item.file_path else None
            ))
        self.db_conn.commit()
    
    def create_knowledge_graph(self, content: List[WebContent]):
        """Create knowledge graph from content"""
        logger.info("Creating knowledge graph...")
        for item in content:
            entities = self.nlp_processor.extract_entities(item.content)
            relationships = self.nlp_processor.extract_relationships(item.content, entities)
            attributes = self.nlp_processor.extract_attributes(item.content, entities)
            for entity in entities:
                self.knowledge_graph.add_entity(entity)
                if entity.text in attributes:
                    self.knowledge_graph.entities[f"{entity.label}_{entity.text}"]['attributes'].update(attributes[entity.text])
            for rel in relationships:
                self.knowledge_graph.add_relationship(rel)
        
        self.knowledge_graph.export_to_json(self.output_dir / 'knowledge_graph.json')
        self.knowledge_graph.visualize(str(self.output_dir / 'knowledge_graph.png'))
        logger.info(f"Graph stats: {self.knowledge_graph.get_statistics()}")
        return self.knowledge_graph
    
    async def run_complete_pipeline(self):
        """Run complete pipeline"""
        logger.info("Starting Bajaj Finserv extraction pipeline...")
        scraped_content = await self.run_data_ingestion()
        knowledge_graph = self.create_knowledge_graph(scraped_content)
        logger.info("Pipeline completed!")
        return scraped_content, knowledge_graph
    
    async def close(self):
        """Close resources"""
        await self.crawler.close()
        self.db_conn.close()

async def main():
    """Main function"""
    extractor = BajajFinservKnowledgeExtractor()
    try:
        scraped_content, knowledge_graph = await extractor.run_complete_pipeline()
        print(f"Processed {len(scraped_content)} pages")
        print(f"Graph: {knowledge_graph.get_statistics()['num_entities']} entities, "
              f"{knowledge_graph.get_statistics()['num_relationships']} relationships")
    finally:
        await extractor.close()

if __name__ == "__main__":
    asyncio.run(main())